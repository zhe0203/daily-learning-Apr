# -*- coding: utf-8 -*-
# http://www.jianshu.com/p/0467fb4c9317
# https://python-docx.readthedocs.io/en/latest/index.html#api-documentation
'''
现在需要通过python对，”表扬信.docx”文档进行修改，需要修改的地方已在图中标记出。
1、第一个箭头处，首行缩进2字符
2、第二个箭头处，对段落进行左缩进2字符，并添加“向小z同学学习！”
3、第三个和第四个箭头处，进行右对齐，并右缩进2cm
4、赵东来，修改为小z
5、陆亦可，修改为大Z
6、她，修改为他
7、狗粮，修改为猫粮
'''

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

# 读取文档
document = Document(r'C:\Users\jk\Desktop\Python学习\每日一学习\docx例子.docx')
# 首先对段落格式进行修改，docx默认标题也属于段落，因此"表扬信"是第一段
paragraphs = document.paragraphs       # 读取段落
# print(paragraphs)
paragraphs[2].paragraph_format.first_line_indent = Cm(0.74)
paragraphs[3].paragraph_format.left_indent = Cm(0.74)
paragraphs[4].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
paragraphs[4].paragraph_format.right_indent = Cm(2)
paragraphs[5].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
paragraphs[5].paragraph_format.right_indent = Cm(2)
# 对文本进行修改
# 修改第二段
paragraphs[1].text = '小z同学:'
# 将第三段陆亦可替换为大z，他替换为她，通过python表达式，可以很简单的进行文本的替换和查找
text = re.sub('陆亦可','大z',paragraphs[2].text)
text = re.sub('他','她',text)
paragraphs[2].text = text
# 在第四段后面加上
paragraphs[3].add_run('向小z同学学习！')
# 修改表格内容
tables = document.tables
tables[0].cell(1,0).text = '猫粮'
tables[0].cell(2,0).text = '猫粮'
tables[0].cell(3,0).text = '猫粮'
# 插入图片
# document.add_picture('xxx.jpg',width=Cm(11.8))
document.save('demo.docx')

# 官方示例
from docx import Document
from docx.shared import Inches

# 建立新的word
document = Document()
# 添加标题
document.add_heading('Document Title',0)

# 添加段落
p = document.add_paragraph('A palin paragraph having some')
# 设置这一段落的格式
p.add_run(' bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading,level 1',level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')
document.add_paragraph('first item in unordered list', style='ListBullet')
document.add_paragraph('first item in ordered list', style='ListNumber')

# 添加表格
table = document.add_table(rows=1,cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'

# for item in recordset:
#     row_cells = tables.add_row().cells
#     row_cells[0].text = str(item.qty)
#     row_cells[1].text = str(item.id)
#     row_cells[2].text = item.desc

document.add_page_break()

document.save('test.docx')
